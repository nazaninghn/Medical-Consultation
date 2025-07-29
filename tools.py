import os
import json
from datetime import datetime
from langchain.tools import Tool
from PIL import Image
import PyPDF2
from docx import Document
import speech_recognition as sr
from gtts import gTTS
import io
import base64

# Save symptom entry to a log file
def log_symptom_entry(entry: str):
    """Log user's symptom consultation with timestamp and structured data"""
    os.makedirs("logs", exist_ok=True)
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    date_str = datetime.now().strftime('%Y%m%d')
    filename = f"logs/symptoms_{date_str}.log"
    
    # Create structured log entry
    log_entry = {
        "timestamp": timestamp,
        "entry": entry,
        "session_type": "symptom_consultation"
    }
    
    with open(filename, "a", encoding="utf-8") as f:
        f.write(f"{json.dumps(log_entry)}\n")
    
    return f"Symptom consultation logged successfully at {timestamp}"

def get_medical_resources():
    """Provide emergency contact information and medical resources"""
    resources = {
        "emergency": "Call 911 for medical emergencies",
        "poison_control": "1-800-222-1222",
        "mental_health": "988 Suicide & Crisis Lifeline",
        "general_advice": "Contact your primary care physician for non-emergency concerns"
    }
    return json.dumps(resources, indent=2)

def process_uploaded_file(file_path: str):
    """Process uploaded files (images, PDFs, documents) and extract relevant information"""
    try:
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']:
            # Process image file
            with Image.open(file_path) as img:
                # Basic image analysis
                return f"Image processed: {img.size[0]}x{img.size[1]} pixels, format: {img.format}"
                
        elif file_extension == '.pdf':
            # Process PDF file
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_content = ""
                for page in pdf_reader.pages:
                    text_content += page.extract_text()
                return f"PDF processed: {len(pdf_reader.pages)} pages, extracted text length: {len(text_content)} characters"
                
        elif file_extension in ['.doc', '.docx']:
            # Process Word document
            doc = Document(file_path)
            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text
            return f"Document processed: {len(doc.paragraphs)} paragraphs, text length: {len(text_content)} characters"
            
        else:
            return f"Unsupported file type: {file_extension}"
            
    except Exception as e:
        return f"Error processing file: {str(e)}"

def convert_speech_to_text(audio_file_path: str):
    """Convert speech audio to text"""
    try:
        recognizer = sr.Recognizer()
        with sr.AudioFile(audio_file_path) as source:
            audio = recognizer.record(source)
            text = recognizer.recognize_google(audio)
            return f"Speech converted to text: {text}"
    except Exception as e:
        return f"Error converting speech to text: {str(e)}"

def convert_text_to_speech(text: str, session_id: str):
    """Convert text to speech and return audio file path"""
    try:
        os.makedirs("static/audio", exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        audio_filename = f"response_{session_id}_{timestamp}.mp3"
        audio_path = f"static/audio/{audio_filename}"
        
        tts = gTTS(text=text, lang='en', slow=False)
        tts.save(audio_path)
        
        return audio_path
    except Exception as e:
        return f"Error converting text to speech: {str(e)}"

tools = [
    Tool(
        name="log_symptom_entry",
        func=log_symptom_entry,
        description="Log user's symptom consultation with timestamp and details"
    ),
    Tool(
        name="get_medical_resources",
        func=get_medical_resources,
        description="Get emergency contacts and medical resource information"
    ),
    Tool(
        name="process_uploaded_file",
        func=process_uploaded_file,
        description="Process uploaded files (images, PDFs, documents) and extract information"
    ),
    Tool(
        name="convert_speech_to_text",
        func=convert_speech_to_text,
        description="Convert speech audio to text"
    ),
    Tool(
        name="convert_text_to_speech",
        func=convert_text_to_speech,
        description="Convert text response to speech audio"
    )
]
